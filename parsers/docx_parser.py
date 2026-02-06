"""
DOCX 文件解析器
使用 python-docx 库提取Word文档的文本和结构。
单次打开文档即可同时得到全文与结构，避免重复 I/O。
"""
import os
from typing import Optional

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _open_doc(file_path: str) -> "Document":
    """打开 DOCX，仅做校验与返回 Document。"""
    if not DOCX_AVAILABLE:
        raise ImportError("请先安装python-docx库: pip install python-docx")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    if not file_path.lower().endswith('.docx'):
        raise ValueError(f"不是DOCX文件: {file_path}")
    try:
        return Document(file_path)
    except PackageNotFoundError:
        raise ValueError(f"无法打开文件，可能不是有效的DOCX文件: {file_path}")
    except Exception as e:
        raise ValueError(f"解析DOCX文件时出错: {e}")


def parse_docx_with_structure(file_path: str) -> tuple[str, list[dict]]:
    """
    一次打开 DOCX，同时返回全文与结构。需要两者时请用此函数以避免重复 I/O。
    Returns:
        (content: str, structure: list[dict])
    """
    doc = _open_doc(file_path)
    paragraphs_text: list[str] = []
    structure: list[dict] = []
    current_section: dict = {"title": "", "level": 0, "content": []}

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            if current_section["title"] or current_section["content"]:
                current_section["content"] = "\n".join(current_section["content"]).strip()
                structure.append(current_section)
            try:
                level = int(style_name.replace("Heading ", "").replace("Heading", "1"))
            except ValueError:
                level = 1
            current_section = {"title": text, "level": level, "content": []}
        else:
            if text:
                current_section["content"].append(text)

        if text:
            paragraphs_text.append(text)

    if current_section["title"] or current_section["content"]:
        current_section["content"] = "\n".join(current_section["content"]).strip()
        structure.append(current_section)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                line = " | ".join(row_text)
                paragraphs_text.append(line)

    content = "\n\n".join(paragraphs_text)
    return content, structure


def parse_docx(file_path: str) -> str:
    """解析 DOCX 文件，返回文本内容。仅需正文时使用。"""
    content, _ = parse_docx_with_structure(file_path)
    return content


def extract_structure(file_path: str) -> list[dict]:
    """从 DOCX 提取文档结构（标题和正文）。仅需结构时使用。"""
    _, structure = parse_docx_with_structure(file_path)
    return structure
