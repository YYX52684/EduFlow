# -*- coding: utf-8 -*-
"""
把老师端模板的 Markdown 生成 docx：
- 支持简单标题层级（# / ## / ###）
- 支持项目符号（- ）和编号列表（1. ）
- 支持基本 Markdown 表格（连续的 | ... | 行）

本脚本只做“尽量保留版式”的转换，目标是让老师可读、结构可被解析器识别。
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


_HEADING_MAP = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _parse_table_rows(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """
    解析 Markdown 表格：
    | a | b |
    | --- | --- |
    | c | d |

    返回 (header, rows)
    """

    def split_cells(row: str) -> list[str]:
        # 去掉首尾 |，再按 | 分割
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        # 过滤掉空单元（避免多余列）
        return [c for c in cells]

    first = split_cells(table_lines[0])
    # 第二行通常是分隔行（由 --- / :---: 组成）
    data_start_idx = 2 if len(table_lines) >= 2 else 1
    rows: list[list[str]] = []
    for line in table_lines[data_start_idx:]:
        if not _is_table_row(line):
            continue
        rows.append(split_cells(line))
    return first, rows


def _add_table(doc: Document, table_lines: list[str]) -> None:
    header, rows = _parse_table_rows(table_lines)
    n_cols = len(header) if header else 1
    n_rows = max(2, len(rows) + 1)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # header
    for j in range(n_cols):
        table.cell(0, j).text = header[j] if j < len(header) else ""

    # data
    for i, r in enumerate(rows, start=1):
        for j in range(n_cols):
            table.cell(i, j).text = r[j] if j < len(r) else ""


def md_to_docx(md_text: str, output_path: Path) -> None:
    doc = Document()

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            doc.add_paragraph(heading_text, style=_HEADING_MAP.get(level, "Normal"))
            i += 1
            continue

        # 表格：连续的 |...| 行
        if _is_table_row(line):
            table_block: list[str] = [line]
            j = i + 1
            while j < len(lines) and _is_table_row(lines[j]):
                table_block.append(lines[j].rstrip())
                j += 1
            _add_table(doc, table_block)
            i = j
            continue

        # 列表：- xxx
        m_bullet = re.match(r"^-\s+(.+)$", line.strip())
        if m_bullet:
            doc.add_paragraph(m_bullet.group(1).strip(), style="List Bullet")
            i += 1
            continue

        # 列表：1. xxx
        m_num = re.match(r"^\d+\.\s+(.+)$", line.strip())
        if m_num:
            doc.add_paragraph(m_num.group(1).strip(), style="List Number")
            i += 1
            continue

        # 其它：普通段落
        doc.add_paragraph(line.strip())
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input-md",
        default="workspaces/u1124360866/input/机械类/机械类能力训练材料参考指南.md",
        help="老师端模板 md 路径",
    )
    parser.add_argument(
        "--output-docx",
        default="workspaces/u1124360866/input/机械类/机械类能力训练材料参考指南.docx",
        help="输出 docx 路径",
    )
    args = parser.parse_args()

    repo_root = Path(__file__).resolve().parents[1]
    input_md = (repo_root / args.input_md).resolve()
    output_docx = (repo_root / args.output_docx).resolve()

    md_text = input_md.read_text(encoding="utf-8")
    md_to_docx(md_text, output_docx)
    print(f"[OK] 已生成：{output_docx}")


if __name__ == "__main__":
    main()

